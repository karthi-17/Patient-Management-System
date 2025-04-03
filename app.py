from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import json
import csv
from io import StringIO, BytesIO
from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl.styles import Font
import sqlite3
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)

# Use environment variable for database URL, fallback to SQLite for local development
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('DATABASE_URL', 'sqlite:///patients_list.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Production configurations
app.config['PREFERRED_URL_SCHEME'] = 'https'

db = SQLAlchemy(app)

class Patient(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    mobile = db.Column(db.String(15), nullable=False)
    visit_date = db.Column(db.String(10), nullable=False)
    prescription = db.Column(db.Text, nullable=False)
    grand_total = db.Column(db.Float, nullable=False)
    discount = db.Column(db.Float, nullable=False, default=0)  # Stores discount percentage
    discount_amount = db.Column(db.Float, nullable=False, default=0)  # Stores calculated discount amount
    total_price = db.Column(db.Float, nullable=False)

    def to_dict(self):
        return {
            'id': self.id,
            'name': self.name,
            'mobile': self.mobile,
            'visit_date': self.visit_date,
            'prescription': json.loads(self.prescription),
            'grand_total': self.grand_total,
            'discount': self.discount,
            'discount_amount': self.discount_amount,
            'total_price': self.total_price
        }

    @property
    def prescription_list(self):
        return json.loads(self.prescription)

with app.app_context():
    db.create_all()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/add_patient')
def add_patient_page():
    return render_template('add_patient.html')

@app.route('/database')
def database():
    return render_template('database.html')

@app.route('/add_patient', methods=['POST'])
def add_patient():
    data = request.json
    prescription_data = json.dumps(data['prescription'])
    
    new_patient = Patient(
        name=data['name'],
        mobile=data['mobile'],
        visit_date=data['visit_date'],
        prescription=prescription_data,
        grand_total=data['grand_total'],
        discount=data['discount'],
        discount_amount=data['discount_amount'],
        total_price=data['total_price']
    )
    
    db.session.add(new_patient)
    db.session.commit()
    return jsonify({'message': 'Patient added successfully'})

@app.route('/search_patients', methods=['GET'])
def search_patients():
    search_type = request.args.get('type')
    search_value = request.args.get('value')
    
    query = Patient.query
    
    if search_type == 'name':
        query = query.filter(Patient.name.like(f'%{search_value}%'))
    elif search_type == 'mobile':
        query = query.filter(Patient.mobile.like(f'%{search_value}%'))
    elif search_type == 'date':
        query = query.filter(Patient.visit_date == search_value)
    
    patients = query.all()
    return jsonify([patient.to_dict() for patient in patients])

@app.route('/get_all_patients')
def get_all_patients():
    patients = Patient.query.all()
    return jsonify([patient.to_dict() for patient in patients])

@app.route('/patient/<int:patient_id>')
def patient_details(patient_id):
    patient = Patient.query.get_or_404(patient_id)
    return render_template('patient_details.html', patient=patient)

@app.route('/delete_patients', methods=['POST'])
def delete_patients():
    try:
        patient_ids = request.json.get('patient_ids', [])
        for patient_id in patient_ids:
            patient = Patient.query.get(patient_id)
            if patient:
                db.session.delete(patient)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})

@app.route('/export_patients')
def export_patients():
    try:
        patient_ids = request.args.get('patient_ids', '')
        export_format = request.args.get('format', 'csv')

        if patient_ids:
            patient_ids = [int(id) for id in patient_ids.split(',')]
            patients = Patient.query.filter(Patient.id.in_(patient_ids)).all()
        else:
            patients = Patient.query.all()

        if export_format == 'csv':
            return export_as_csv(patients)
        elif export_format == 'xlsx':
            return export_as_excel(patients)
        elif export_format == 'docx':
            return export_as_word(patients)
        elif export_format == 'db':
            return export_as_db(patients)
        else:
            return jsonify({'success': False, 'error': 'Invalid export format'})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

def export_as_csv(patients):
    si = StringIO()
    cw = csv.writer(si)
    cw.writerow(['Name', 'Mobile', 'Visit Date', 'Prescription'])
    
    for patient in patients:
        cw.writerow([
            patient.name,
            patient.mobile,
            patient.visit_date,
            patient.prescription
        ])
    
    output = si.getvalue()
    si.close()
    
    return send_file(
        StringIO(output),
        mimetype='text/csv',
        as_attachment=True,
        download_name='patients.csv'
    )

def export_as_excel(patients):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Patients"

    # Add headers
    headers = ['Name', 'Mobile', 'Visit Date', 'Prescription']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col)
        cell.value = header
        cell.font = Font(bold=True)

    # Add data
    for row, patient in enumerate(patients, 2):
        ws.cell(row=row, column=1).value = patient.name
        ws.cell(row=row, column=2).value = patient.mobile
        ws.cell(row=row, column=3).value = patient.visit_date
        ws.cell(row=row, column=4).value = patient.prescription

    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width

    excel_file = BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)

    return send_file(
        excel_file,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='patients.xlsx'
    )

def export_as_word(patients):
    doc = Document()
    doc.add_heading('Patient Records', 0)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['Name', 'Mobile', 'Visit Date', 'Prescription']
    for i, header in enumerate(headers):
        header_cells[i].text = header

    for patient in patients:
        row_cells = table.add_row().cells
        row_cells[0].text = patient.name
        row_cells[1].text = patient.mobile
        row_cells[2].text = patient.visit_date
        row_cells[3].text = patient.prescription

    doc_file = BytesIO()
    doc.save(doc_file)
    doc_file.seek(0)

    return send_file(
        doc_file,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='patients.docx'
    )

def export_as_db(patients):
    db_file = BytesIO()
    
    # Create a new SQLite database in memory
    conn = sqlite3.connect(':memory:')
    c = conn.cursor()
    
    # Create the patients table
    c.execute('''CREATE TABLE patients
                 (name text, mobile text, visit_date text, prescription text)''')
    
    # Insert the data
    for patient in patients:
        c.execute("INSERT INTO patients VALUES (?, ?, ?, ?)",
                 (patient.name, patient.mobile, patient.visit_date, patient.prescription))
    
    conn.commit()
    
    # Save the in-memory database to a file
    temp_db = BytesIO()
    for data in conn.iterdump():
        temp_db.write(data.encode('utf-8'))
    temp_db.seek(0)
    
    conn.close()
    
    return send_file(
        temp_db,
        mimetype='application/x-sqlite3',
        as_attachment=True,
        download_name='patients.db'
    )

@app.route('/import_patients', methods=['POST'])
def import_patients():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        
        if not file.filename.endswith('.csv'):
            return jsonify({'success': False, 'error': 'Only CSV files are allowed'})
        
        stream = StringIO(file.stream.read().decode("UTF8"), newline=None)
        csv_input = csv.reader(stream)
        next(csv_input)  # Skip header row
        
        for row in csv_input:
            if len(row) >= 4:
                patient = Patient(
                    name=row[0],
                    mobile=row[1],
                    visit_date=row[2],
                    prescription=row[3]
                )
                db.session.add(patient)
        
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    # Only enable debug mode in development
    debug_mode = os.getenv('FLASK_ENV') == 'development'
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', 5000)), debug=debug_mode)
